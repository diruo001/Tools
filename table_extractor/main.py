import argparse

import bs4
from docx import Document
from docx.table import _Cell
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import numpy as np
import pandas as pd
import os

from html_parser import Table, extract_tables


def save_table_in_doc(doc, table: Table):
    if table.caption is not None:
        table_doc = doc.add_table(rows=1, cols=1, style="Table Grid")
        paragraph = table_doc.cell(0, 0).paragraphs[0]
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for index, content in enumerate(table.caption):
            if isinstance(content, str):
                if index == 0:
                    content = content.lstrip()
                run = paragraph.add_run(content.replace("\n", ""))
                run.font.name = 'Times New Roman'
            if isinstance(content, bs4.element.Tag):
                run = paragraph.add_run(content.contents[0].replace("\n", ""))
                run.font.name = 'Times New Roman'
                if content.name == "sub":
                    run.font.subscript = True
                if content.name == "sup":
                    run.font.superscript = True
    rows, cols = table.shape()
    if table.headers is not None and cols > 0:
        table_doc = doc.add_table(rows=1, cols=cols, style="Table Grid")
        for j in range(cols):
            paragraph = table_doc.cell(0, j).paragraphs[0]
            for index, content in enumerate(table.headers[j]):
                if isinstance(content, str):
                    if index == 0:
                        content = content.lstrip()
                    run = paragraph.add_run(content.replace("\n", ""))
                    run.font.name = 'Times New Roman'
                if isinstance(content, bs4.element.Tag):
                    run = paragraph.add_run(content.contents[0].replace("\n", ""))
                    run.font.name = 'Times New Roman'
                    if content.name == "sub":
                        run.font.subscript = True
                    if content.name == "sup":
                        run.font.superscript = True
    table_doc = doc.add_table(rows=rows, cols=cols, style="Table Grid")
    for i in range(rows):
        for j in range(cols):
            paragraph = table_doc.cell(i, j).paragraphs[0]
            for index, content in enumerate(table.rows[i][j]):
                if isinstance(content, str):
                    if index == 0:
                        content = content.lstrip()
                    run = paragraph.add_run(content.replace("\n", ""))
                    run.font.name = 'Times New Roman'
                if isinstance(content, bs4.element.Tag):
                    run = paragraph.add_run(content.contents[0].replace("\n", ""))
                    run.font.name = 'Times New Roman'
                    if content.name == "sub":
                        run.font.subscript = True
                    if content.name == "sup":
                        run.font.superscript = True
    doc.add_page_break()


def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
 
    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
 
            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
 
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--data_dir", type=str, default="test_data")
    parser.add_argument("--table_idx_list", type=int, nargs="+", default=1)
    parser.add_argument("--save_path", type=str, default="result.docx")
    args = parser.parse_args()

    if not os.path.exists(args.data_dir):
        raise f"{args.data_dir} is not exist."

    document = Document()
    table_names = []
    for idx in args.table_idx_list:
        table_names.append(f"Table {idx}")
    print("extract:", table_names)

    for filename in os.listdir(args.data_dir):
        document.add_paragraph()
        print(filename)
        filepath = os.path.join(args.data_dir, filename)
        table_list = extract_tables(filepath, table_names)
        for table in table_list:
            save_table_in_doc(doc=document, table=table)
    document.save(args.save_path)